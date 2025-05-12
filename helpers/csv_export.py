import pandas as pd
import os
# import textract
import PyPDF2
from docx import Document
import logging
import tempfile
import shutil
from zipfile import ZipFile, BadZipFile
import csv
import openpyxl
import xlrd
import subprocess
# import win32com.client



# Налаштування логування
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def convert_to_csv(input_path, output_path=None):
    """
    Конвертує файли різних форматів (.docx, .doc, .xls, .xlsx, pdf) у CSV формат.
    
    Args:
        input_path: Шлях до вхідного файлу
        output_path: Шлях до вихідного CSV файлу. Якщо не вказано, 
                     буде створено файл з тим самим ім'ям, але з розширенням .csv
    
    Returns:
        Шлях до створеного CSV файлу
    
    Raises:
        ValueError: Якщо формат файлу не підтримується
    """
    ext = os.path.splitext(input_path)[-1].lower()

    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.csv'
    
    logger.info(f"Конвертація файлу {input_path} у формат CSV: {output_path}")

    try:
        if ext in ['.xlsx', '.xls']:
            _convert_excel_to_csv(input_path, output_path, ext)
        elif ext == '.docx':
            _convert_docx_to_csv(input_path, output_path)
        elif ext == '.doc':
            # convert_doc_to_csv(input_path, output_path)
            raise ValueError(f"Формат {ext} не підтримується!")
        elif ext == '.pdf':
            _convert_pdf_to_csv(input_path, output_path)
        else:
            raise ValueError(f"Формат {ext} не підтримується!")
        
        logger.info(f"Конвертація завершена успішно. Створено файл: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Помилка при конвертації файлу {input_path}: {str(e)}")
        raise

def _convert_excel_to_csv(input_path, output_path, ext):
    """
    Конвертує Excel файли (.xlsx, .xls) у CSV формат з використанням різних методів.
    Спробує кілька підходів, якщо один з них не спрацює.
    
    Args:
        input_path: Шлях до вхідного Excel файлу
        output_path: Шлях до вихідного CSV файлу
        ext: Розширення файлу (.xlsx або .xls)
    """
    methods = [
        _convert_excel_pandas,
        _convert_excel_openpyxl if ext == '.xlsx' else _convert_excel_xlrd,
        _convert_excel_manual
    ]
    
    last_error = None
    for method in methods:
        try:
            logger.info(f"Спроба конвертації Excel файлу методом {method.__name__}")
            method(input_path, output_path)
            logger.info(f"Конвертація успішна методом {method.__name__}")
            return
        except Exception as e:
            logger.warning(f"Метод {method.__name__} не спрацював: {str(e)}")
            last_error = e
    
    # Якщо жоден метод не спрацював
    raise Exception(f"Не вдалося конвертувати Excel файл. Остання помилка: {str(last_error)}")

def _convert_excel_pandas(input_path, output_path):
    """Конвертація Excel файлу за допомогою pandas"""
    df = pd.read_excel(input_path, engine='openpyxl' if input_path.endswith('.xlsx') else 'xlrd')
    df.to_csv(output_path, index=False, encoding='utf-8')

def _convert_excel_openpyxl(input_path, output_path):
    """Конвертація XLSX файлу за допомогою openpyxl"""
    wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
    sheet = wb.active
    
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        for row in sheet.rows:
            writer.writerow([cell.value for cell in row])

def _convert_excel_xlrd(input_path, output_path):
    """Конвертація XLS файлу за допомогою xlrd"""
    book = xlrd.open_workbook(input_path)
    sheet = book.sheet_by_index(0)
    
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        for row_idx in range(sheet.nrows):
            writer.writerow([sheet.cell_value(row_idx, col_idx) for col_idx in range(sheet.ncols)])

def _convert_excel_manual(input_path, output_path):
    """
    Ручна конвертація Excel файлу шляхом розпакування XLSX як ZIP архіву.
    Цей метод використовується як запасний варіант, якщо інші методи не спрацювали.
    """
    # Створюємо тимчасову копію файлу
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(input_path)[1]) as tmp:
        shutil.copy2(input_path, tmp.name)
        tmp_path = tmp.name
    
    try:
        # Спроба розпакувати XLSX як ZIP архів
        with ZipFile(tmp_path) as zf:
            # Виводимо вміст архіву для діагностики
            file_list = zf.namelist()
            logger.info(f"Вміст архіву: {', '.join(file_list[:10])}...")
            
            # Шукаємо файли worksheet
            worksheet_files = [f for f in file_list if f.startswith('xl/worksheets/sheet')]
            if not worksheet_files:
                raise ValueError("Не знайдено файлів worksheet в архіві")
            
            # Беремо перший лист
            worksheet_file = worksheet_files[0]
            
            # Витягуємо дані з листа
            import xml.etree.ElementTree as ET
            
            # Читаємо XML файл worksheet
            ws_xml = zf.read(worksheet_file)
            ws_root = ET.fromstring(ws_xml)
            
            # Знаходимо всі рядки
            rows = ws_root.findall('.//{*}row')
            logger.info(f"Знайдено {len(rows)} рядків у worksheet")
            
            # Словник для зберігання shared strings, якщо вони є
            shared_strings = []
            
            # Спробуємо знайти shared strings (з урахуванням регістру)
            shared_strings_path = None
            for path in file_list:
                if path.lower() == 'xl/sharedstrings.xml' or path.lower() == 'xl/sharedstrings.xml':
                    shared_strings_path = path
                    break
            
            if shared_strings_path:
                try:
                    logger.info(f"Читання shared strings з {shared_strings_path}")
                    ss_xml = zf.read(shared_strings_path)
                    ss_root = ET.fromstring(ss_xml)
                    
                    # Знаходимо всі елементи si/t
                    for si in ss_root.findall('.//{*}si'):
                        texts = []
                        for t in si.findall('.//{*}t'):
                            if t.text:
                                texts.append(t.text)
                        shared_strings.append(''.join(texts))
                    
                    logger.info(f"Прочитано {len(shared_strings)} shared strings")
                except Exception as e:
                    logger.warning(f"Помилка при читанні shared strings: {str(e)}")
            
            # Відкриваємо CSV файл для запису
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Обробляємо кожен рядок
                for row in rows:
                    csv_row = []
                    
                    # Знаходимо всі комірки в рядку
                    cells = row.findall('.//{*}c')
                    
                    for cell in cells:
                        # Отримуємо значення комірки
                        value = None
                        v_element = cell.find('.//{*}v')
                        
                        if v_element is not None and v_element.text:
                            # Перевіряємо тип комірки
                            cell_type = cell.get('t')
                            
                            if cell_type == 's' and shared_strings:
                                # Це індекс у shared strings
                                try:
                                    idx = int(v_element.text)
                                    if idx < len(shared_strings):
                                        value = shared_strings[idx]
                                    else:
                                        value = v_element.text
                                except (ValueError, IndexError):
                                    value = v_element.text
                            else:
                                # Звичайне значення
                                value = v_element.text
                        
                        csv_row.append(value if value is not None else "")
                    
                    # Записуємо рядок у CSV
                    writer.writerow(csv_row)
                
                logger.info(f"Записано {len(rows)} рядків у CSV файл")
    except Exception as e:
        logger.error(f"Помилка при ручній конвертації Excel файлу: {str(e)}")
        # Створюємо базовий CSV, щоб не повернути порожній файл
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["Помилка при конвертації"])
            writer.writerow([str(e)])
    finally:
        # Видаляємо тимчасовий файл
        try:
            os.unlink(tmp_path)
        except Exception as e:
            logger.warning(f"Не вдалося видалити тимчасовий файл: {str(e)}")

def _convert_docx_to_csv(input_path, output_path):
    """Конвертує DOCX файл у CSV формат"""
    doc = Document(input_path)
    
    # Спочатку перевіряємо наявність таблиць
    if doc.tables:
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for table in doc.tables:
                for row in table.rows:
                    writer.writerow([cell.text.strip() for cell in row.cells])
    else:
        # Якщо таблиць немає, записуємо текст параграфів
        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["Text"])
            for line in text.split('\n'):
                if line.strip():
                    writer.writerow([line.strip()])

# def convert_doc_to_csv(input_path, output_path):
#     """Зчитує таблиці з .doc (або .docx) і експортує у CSV"""
#     word = win32com.client.Dispatch("Word.Application")
#     word.Visible = False  # не показуємо вікно Word

#     try:
#         doc = word.Documents.Open(doc_path)
#         tables = doc.Tables

#         if tables.Count == 0:
#             print("❌ У документі немає таблиць.")
#             return

#         with open(csv_path, 'w', newline='', encoding='utf-8') as f:
#             writer = csv.writer(f)

#             for t_index, table in enumerate(tables):
#                 writer.writerow([f"=== Таблиця {t_index + 1} ==="])

#                 for row in range(1, table.Rows.Count + 1):
#                     row_data = []
#                     for col in range(1, table.Columns.Count + 1):
#                         cell_text = table.Cell(row, col).Range.Text
#                         clean_text = cell_text.replace('\r\x07', '').strip()
#                         row_data.append(clean_text)
#                     writer.writerow(row_data)

#         print(f"✅ Таблиці збережено у CSV: {csv_path}")

#     except Exception as e:
#         print(f"❌ Помилка при обробці файлу: {e}")

#     finally:
#         doc.Close(False)
#         word.Quit()

def _convert_pdf_to_csv(input_path, output_path):
    """Конвертує PDF файл у CSV формат"""
    try:
        text = ""
        with open(input_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        # Записуємо текст у CSV
        lines = text.split('\n')
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f, quoting=csv.QUOTE_ALL, escapechar='\\', quotechar='"', doublequote=True)
            writer.writerow(["Text"])
            for line in lines:
                if line.strip():
                    writer.writerow([line.strip()])
    except Exception as e:
        logger.error(f"Помилка при обробці PDF файлу: {str(e)}")
        raise
